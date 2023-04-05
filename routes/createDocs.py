from fastapi import APIRouter,status, Response,BackgroundTasks, Depends
from fastapi.responses import FileResponse

from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import random
import string

# from db.client import db_inmuebles, db_asociaciones
# from schemas.inmuebles import inmuebleEntity, inmueblesEntity
# from schemas.asociaciones import asociacionEntity,asociacionesEntity
# from models.inmuebles import InmuebleModel
# from models.asociaciones import AsociacioneModels
from bson import ObjectId
from fastapi.middleware.cors import CORSMiddleware
# from starlette.responses import FileResponse,StreamingResponse

# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter, A4

# from reportlab.lib.styles import ParagraphStyle
# from reportlab.lib.styles import getSampleStyleSheet

from auth.auth_handler import signJWT
from auth.auth_bearer import JWTBearer

import pymysql.cursors
from dotenv import dotenv_values,load_dotenv
import os   
import json
import io
import ftplib
import datetime

import db.ConnToMysql as dataBase

# load_dotenv('.env') 

def id_generator(size=6, chars=string.ascii_uppercase + string.digits):
    return ''.join(random.choice(chars) for _ in range(size))
nameAleatorio = id_generator()

# from fastapi.encoders import jsonable_encoder
# from fastapi.responses import JSONResponse
# from fastapi.middleware.cors import CORSMiddleware
# from deta import Deta
# from db import client
# import json

### Inmuebles API ###

# Levantar el server: uvicorn main:app --reload
# Detener el server: CTRL+C

# Documentación con Swagger: http://127.0.0.1:8000/docs
# Documentación con Redocly: http://127.0.0.1:8000/redoc

# DETA INSTRUCTIONS
# deta visor open //para abrir la consola
# deta watch // deploy automaticamente los cambios
# deta --help 

docs = APIRouter()

# deta = Deta() 

# db = deta.Base('inmuebles') #Nombrepara la bbdd

# app = FastAPI()

year = datetime.date.today().year



@docs.get("/arras/{id}")
async def docs_arras(id:int):
    db = dataBase.get_inmueble_id(id)
    dbVendedores = dataBase.get_vendedor_id_inmuebles(id)
    dbCompradores = dataBase.get_comprador_id_inmuebles(id)
    if dbVendedores:
        db['vendedores'] = dbVendedores
    else:
        db['vendedores'] = []
    
    if dbCompradores:
        db['compradores'] = dbCompradores
    else:
        db['compradores'] = []
    
    print('db',db)
    #try:
        #connection = pymysql.connect(
        #host=os.environ.get("hostDB"),
        #user=os.environ.get("userDB"),
        #password=os.environ.get("passwordDB"),
        #database=os.environ.get("databaseDB"),
        #cursorclass=pymysql.cursors.DictCursor)

        #with connection.cursor() as cursor:
            
            #query = f"SELECT * FROM avap.inmuebles WHERE id = {id}"
            #cursor.execute(query)
            #db = cursor.fetchone()
            #print("Resultados de db: ", db)
    tipologia = db['tipologia']
    direccion = db['direccion']
    vendedores = db['vendedores']
    compradores = db['compradores']
    refCatastral = db['refCatastral']
    superficie = db['superficie']
    inscripcionRegistro = db['inscripcionRegistro']
    cru = db['cru']
    precio = db['precio']
    condiPrestamo = db['condiPrestamo']
    
    
    document = Document()

    document.add_heading(f'CONTRATO DE ARRAS PENITENCIALES', 0).center=True
    
    fecha = document.add_paragraph()
    fecha_format = fecha.paragraph_format
    fecha_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha.add_run(f'En _________, a __ de ________ de 202__')
    
    g=document.add_paragraph()
    g_format = g.paragraph_format
    g_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    g.add_run(f'REUNIDOS').bold = True
    # g.add_run(f'REUNIDOS,').bold = True
    e = document.add_paragraph()
    e.add_run(f'De una parte,').bold = True
    for vendedor in vendedores:
        nameV = vendedor['nombre']
        estadoCivilV = vendedor['estadoCivil']
        dniV = vendedor['dni']
        direccionV = vendedor['direccion']
        municipioV = vendedor['municipio']
        provinciaV = vendedor['provincia']
        
        p = document.add_paragraph('D/Dª ')
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f'{nameV}, ').bold = True
        p.add_run('mayor de edad, estado civil ')
        p.add_run(f'{estadoCivilV}, ')
        p.add_run(f'con DNI/NIE ')
        p.add_run(f'{dniV}, ')
        p.add_run(f'con dirección en ')
        p.add_run(f'{direccionV}, ')
        p.add_run(f'de ')
        p.add_run(f'{municipioV} ')
        p.add_run(f'({provinciaV})')
        
    b=document.add_paragraph('Quienes actúan como parte ')
    b.add_run(f'VENDEDORA.').bold = True
    d = document.add_paragraph()
    d.add_run(f'Y, de otra parte,').bold = True
    for comprador in compradores:
        nameC = comprador['nombre']
        estadoCivilC = comprador['estadoCivil']
        dniC = comprador['dni']
        direccionC = comprador['direccion']
        municipioC = comprador['municipio']
        provinciaC = comprador['provincia']

        

        
        c = document.add_paragraph('D/Dª ')
        c_format = c.paragraph_format
        c_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run(f'{nameC}, ').bold = True
        c.add_run('mayor de edad, estado civil ')
        c.add_run(f'{estadoCivilC}, ')
        c.add_run(f'con DNI/NIE ')
        c.add_run(f'{dniC}, ')
        c.add_run(f'con dirección en ')
        c.add_run(f'{direccionC}, ')
        c.add_run(f'de ')
        c.add_run(f'{municipioC} ')
        c.add_run(f'({provinciaC})')
    
    f=document.add_paragraph('Quienes actúan como parte ')
    f.add_run(f'COMPRADORA.').bold = True

    h=document.add_paragraph()
    h_format = h.paragraph_format
    h_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run(f'INTERVIENEN').bold = True

    c = document.add_paragraph()
    c_format = c.paragraph_format
    c_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c.add_run('Cada uno de ellos en su propio nombre y derecho y se reconocen mutuamente capacidad legal suficiente para el otorgamiento del presente ')
    c.add_run(f'CONTRATO DE ARRAS O SEÑAL ').bold = True
    c.add_run(', a cuyo efecto;')

    i=document.add_paragraph()
    i_format = i.paragraph_format
    i_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(f'MANIFESTACIONES').bold = True

    j = document.add_paragraph()
    j_format = j.paragraph_format
    j_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    j.add_run(f'I.- ').bold = True
    j.add_run(f'Que la parte VENDEDORA es propietaria en pleno dominio de un inmueble sito en {direccion}, número de Referencia Catastral: {refCatastral}, con una superficie aproximada de {superficie} y está inscrita en el Registro de la Propiedad de {inscripcionRegistro}. Código Registral Único: {cru}')
    
    k = document.add_paragraph()
    k_format = k.paragraph_format
    k_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    k.add_run(f'II.- ').bold = True
    k.add_run(f'Que la parte VENDEDORA está interesada en vender y la parte COMPRADORA está interesada en comprar con carácter privativo, en el estado actual y en las condiciones en que se halla el inmueble reseñado en la manifestación I. Declara la parte compradora que conoce el estado del inmueble y sus anejos inseparables si los tuviese y lo acepta en su estado actual como cuerpo cierto.')

    l = document.add_paragraph('Con base a lo anterior, deciden formalizar un CONTRATO DE ARRAS y con el fin de regularlo pactan las siguientes;')
    l_format = l.paragraph_format
    l_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    m=document.add_paragraph()
    m_format = m.paragraph_format
    m_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(f'ESTIPULACIONES').bold = True
    
    n = document.add_paragraph()
    n_format = n.paragraph_format
    n_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    n.add_run(f'PRIMERA.- ').bold = True
    n.add_run(f'El precio de la compraventa se fija en ')
    n.add_run(f'XXXXXXXX EUROS ({precio} €) ').bold = True
    n.add_run(f'más los impuestos y gastos que correspondan.')

    o = document.add_paragraph()
    o_format = o.paragraph_format
    o_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    o.add_run(f'SEGUNDA.- ').bold = True
    o.add_run(f'La forma de pago del precio estipulado será:')

    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f'A.- ').bold = True
    p.add_run(f'Un primer pago que realiza la parte COMPRADORA por la cantidad de ')
    p.add_run(f' ______ EUROS (00.000,00 €) ').bold = True
    p.add_run(f'realizado mediante transferencia bancaria de fecha _______ a la cuenta número: ES____ de __. Perteneciente a la parte VENDEDORA, y cuyo justificante se aporta al presente contrato de arras.')

    q = document.add_paragraph()
    q_format = q.paragraph_format
    q_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    q.add_run(f'B.- ').bold = True
    q.add_run(f'Si en el transcurso de los próximos tres días la parte VENDEDORA no hubiera recibido en su cuenta la transferencia de ')
    q.add_run(f' ______ EUROS (00.000,00 €) ').bold = True
    q.add_run(f' indicado en la estipulación A, el presente contrato de Arras quedará nulo y sin efecto alguno.')

    r = document.add_paragraph()
    r_format = r.paragraph_format
    r_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r.add_run(f'C.- ').bold = True
    r.add_run(f'Que la cantidad expresada en la estipulación A, una vez verificada por la parte VENDEDORA, se descontará del precio total de la futura compraventa que se formalice, quedando la misma en poder de la parte VENDEDORA.')

    s = document.add_paragraph()
    s_format = s.paragraph_format
    s_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.add_run(f'D.- ').bold = True
    s.add_run(f'La cantidad restante que asciende a _____________ MIL EUROS (________,__ €) más los impuestos y gastos que correspondan, se abonará en el momento en el que se formalice la escritura de compraventa ante notario.')

    t = document.add_paragraph()
    t_format = t.paragraph_format
    t_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    t.add_run(f'TERCERA.- ').bold = True
    t.add_run(f'La transmisión de la propiedad y la posesión de la finca descrita, se producirá en el momento del otorgamiento de la correspondiente escritura pública de compraventa y completo pago del precio pactado. Con el cumplimiento de los citados requisitos la parte COMPRADORA adquirirá la propiedad del inmueble, momento en el cual tomará en concepto de dueño la posesión material y directa del mismo.')

    u = document.add_paragraph()
    u_format = u.paragraph_format
    u_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    u.add_run(f'CUARTA.- ').bold = True
    u.add_run(f'El otorgamiento de la correspondiente escritura pública de compraventa se realizará antes del _____________________ como fecha máxima, en la notaría que designe la parte COMPRADORA, fijándose la fecha de escrituración de mutuo acuerdo entre ambas partes antes de los últimos 15 días de la fecha máxima del presente contrato.')

    v = document.add_paragraph()
    v_format = v.paragraph_format
    v_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    v.add_run(f'Todo lo cual deberá comunicarse a las partes con una antelación mínima de 10 días. La entrega de llaves y toma de posesión se realizará el día de la firma en notaría, de no ser entregadas por la parte vendedora, esta será penalizada con la cantidad diaria de CIEN EUROS (100,00€).')

    w = document.add_paragraph()
    # w.style = document.styles.add_style('Style name', WD_STYLE_TYPE.PARAGRAPH)
    # font = w.style.font
    # font.name ='Times New Roman'
    # font.size = Pt(9)
    # font.color.rgb = RGBColor(247, 6, 6)
    # document.add_page_break()
    w_format = w.paragraph_format
    w_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    w.add_run(f'Dado el carácter penitencial de estas arras y de acuerdo con lo establecido en el artículo 1.454 del Código Civil, si llegada la fecha máxima para formalizar la escritura de la compraventa la parte COMPRADORA hubiera incumplido lo convenido en el presente contrato, perderá la cantidad de ____________ EUROS (_____,__ €) entregados a cuenta, y no tendrán ningún derecho de compra sobre el inmueble objeto del presente documento. ')
    if condiPrestamo:
        w.add_run(f'Excepto en el caso que a la parte COMPRADORA no le fuera concedido el préstamo hipotecario; en este caso la parte VENDEDORA devolverá el valor de arras íntegro a la parte COMPRADORA.')

        x = document.add_paragraph()
        x_format = x.paragraph_format
        x_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        x.add_run(f'No obstante a lo anterior, dicha excepción (esto es, el derecho al reintegro de las cantidad entregada como señal sin penalización) sólo operará en el caso de que la parte COMPRADORA acredite que la imposibilidad de obtener la financiación necesaria se debe a causas totalmente ajenas a su voluntad. A tal efecto, la parte COMPRADORA deberá aportar a AVAP Agencia de la Vivienda SL (la AGENCIA) la documentación justificativa de dicho extremo.')

    y = document.add_paragraph()
    y_format = y.paragraph_format
    y_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    y.add_run(f'Para el caso de que el préstamo esté aprobado y la operación se pueda firmar en un plazo máximo de diez días después del vencimiento del presente contrato, siempre que sólo sea un problema de fecha de firma en Notaría, las partes podrán pactar la ampliación del presente contrato señalando el día de firma en Notaría con posterioridad al vencimiento del presente documento.')

    z = document.add_paragraph()
    z_format = z.paragraph_format
    z_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    z.add_run(f'Si fuera la parte VENDEDORA quien incumpliera lo convenido en el presente contrato, deberá abonar a la parte COMPRADORA el doble de la cantidad recibida en concepto de arras o señal, es decir, la suma total de ____________ EUROS (_____,__ €)')

    za = document.add_paragraph()
    za_format = za.paragraph_format
    za_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    za.add_run(f'QUINTA.- La finca indicada en las manifestaciones I, se transmitirá libre de cargas y gravámenes, arrendamientos y ocupantes y al corriente de pago de la comunidad de propietarios e impuestos (a la escritura se aportarán los certificados correspondientes).').bold = True

    zb = document.add_paragraph()
    zb_format = zb.paragraph_format
    zb_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    zb.add_run(f'Todos los gastos e impuestos y derechos de este contrato, y de la escritura pública de compraventa que en su día se otorgue (dentro del plazo fijado en la estipulación cuarta), serán de cuenta de la parte COMPRADORA, excepto el pago del Impuesto sobre el Incremento del Valor de los Terrenos de Naturaleza Urbana (plusvalía municipal) que corresponderá a la parte VENDEDORA.')

    zc = document.add_paragraph()
    zc_format = zc.paragraph_format
    zc_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    zc.add_run(f'Las partes pactan que se prorrateará el total del IBI de este año {year} perteneciente los gastos anteriores a la firma a la parte VENDEDORA y en el momento de la firma en adelante será por cuenta de la parte COMPRADORA.')

    zd = document.add_paragraph()
    zd_format = zd.paragraph_format
    zd_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    zd.add_run(f'SÉPTIMA.- ').bold = True
    zd.add_run(f'Las partes se comprometen en todo momento a cumplir el presente contrato y a ejercitar los derechos y cumplir las obligaciones que se desprenden del mismo, conforme a las más estrictas exigencias de la buena fe.')

    ze = document.add_paragraph()
    ze_format = ze.paragraph_format
    ze_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ze.add_run(f'OCTAVA.- ').bold = True
    ze.add_run(f'A efectos de notificaciones y requerimientos, las partes señalan como domicilios los designados al principio de este contrato como propios de los intervinientes.')

    zf = document.add_paragraph()
    zf_format = zf.paragraph_format
    zf_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    zf.add_run(f'NOVENA.- ').bold = True
    zf.add_run(f'Para cualquiera incidencia o conflicto derivado del presente contrato, de su cumplimiento, ejecución o rescisión, las partes con expresa renuncia de su fuero propio, se someten a los Juzgados y Tribunales de _____________________.')

    document.add_page_break()

    zg = document.add_paragraph()
    zg_format = zg.paragraph_format
    zg_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    zg.add_run(f'Y para que así conste, una vez leído, firman este contrato por duplicado y a un solo efecto, en el lugar y fecha indicados en el encabezamiento')

    numVende = len(vendedores)
    datos = []
    for vendedor in vendedores:
        nameV = vendedor['nombre']
        dniV = vendedor['dni']
        vende = nameV, dniV
        print('vende',)
        datos.append(vende)
    print('ooooooooo', datos)
    datos = tuple(datos)
    print('aaaaaaa', datos)

    records = datos

    # table = document.add_table(rows=numVende-1, cols=2)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'PARTE VENDEDORA'
    # hdr_cells[1].text = 'PARTE COMPRADORA'
    # for qty, id in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells = table.add_row().cells
    #     row_cells[1].text = id






        # row_cells[2].text = desc
        # p.add_run('italic.').italic = True

        # p = document.add_paragraph(f'D/Dª {nameV}').bold = True
        # c.add_run(vendedor['nombre'])
        # p.add_run(' and some ')
        
        # c.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph('first item in unordered list', style='List Bullet')
    # document.add_paragraph('first item in ordered list', style='List Number')

            # document.add_picture('monty-truth.png', width=Inches(1.25))

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam'))

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    document.add_page_break()

    document.save(f'files/arras_{nameAleatorio}.docx')
    # threFile = document.save('demo.docx')

    bio = io.BytesIO()
    document.save(bio)  # save to memory stream
    bio.seek(0)  # rewind the stream
            

            ### FTPPPPPPP --------------------------------------------

    # domain name or server ip:
    ftpHost = os.environ.get("ftp_host")
    ftpPort = 21
    ftpUname = os.environ.get("ftp_user")
    ftpPass = os.environ.get("ftp_password")

    ftp = ftplib.FTP(timeout=30)
    ftp.connect(ftpHost, ftpPort)
    ftp.login(ftpUname, ftpPass)

    # fnames = ftp.nlst()
    # ftp.cwd("/var/www/html/backPyAvap/files")
    ftp.cwd("/var/www/html/frontPyAvap/files")

    localFilePath = f'files/arras_{nameAleatorio}.docx'

    with open(localFilePath, 'rb') as file:
        print('ffff',file)
        retCode =ftp.storbinary(f'STOR arras_{nameAleatorio}.docx', file, blocksize=1024*1024)

        ftp.quit()

        if retCode.startswith('256'):
            print('upload success')
        else:
            print('upload not success...') 

        print('Ejecucion completa')

            ### FTPPPPPPP --------------------------------------------
            
        # return f'http://api.acapagencia.com/files/{localFilePath}' 
        # return FileResponse(bio, media_type='application/octet-stream')
        # return StreamingResponse(bio.read(), media_type='application/octet-stream')
    response = {
            "headers": {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Credentials': True
            },
            "statusCode": 200,
            'body': json.dumps({'status':'susccess', 'results':f'Se ha generado el contrato de arras y puedes descargarlo <a href="http://panel.avapagencia.com/{localFilePath}">aquí</a>'})
            }
    return response
    
    # return f'http://panel.avapagencia.com/files/{localFilePath}'

        